"""Document conversion logic for CentralDepo workflow.

Converts company files (docx, doc, xls) to Markdown format.
PDF files are converted using Mistral OCR via the workflows plugin,
passing documents as base64 data URIs.
"""

import asyncio
import base64
import logging
import shutil
import subprocess
from pathlib import Path

import docx2txt
import xlrd
from docx import Document
from mistralai.client.models import DocumentURLChunk
from mistralai.workflows.plugins.mistralai import OCRRequest, mistralai_ocr

from .config import MAX_CONCURRENT_CONVERSIONS, MAX_PDF_SIZE_BYTES
from .downloader import get_company_folder_name

logger = logging.getLogger(__name__)


def _table_to_md(table) -> str:
    """Convert a python-docx table to markdown format.

    Args:
        table: A python-docx table object

    Returns:
        Markdown table string
    """
    md_lines = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        if i == 0:
            # Header row
            md_lines.append("|" + "|".join(cells) + "|")
            md_lines.append("|" + "|".join("-" for _ in cells) + "|")
        else:
            # Data row
            md_lines.append("|" + "|".join(cells) + "|")
    return "\n".join(md_lines)


def _extract_docx(file_path: Path) -> str:
    """Extract text and tables from .docx file.

    Args:
        file_path: Path to the .docx file

    Returns:
        Markdown content string
    """
    doc = Document(str(file_path))
    md_content = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            md_content.append(text)

    for table in doc.tables:
        md_table = _table_to_md(table)
        if md_table:
            md_content.append("")
            md_content.append(md_table)
            md_content.append("")

    return "\n\n".join(md_content)


def _extract_doc(file_path: Path) -> str:
    """Extract text from .doc file using multiple methods.

    Tries multiple approaches in order:
    1. python-docx (for .docx files mislabeled as .doc)
    2. docx2txt (for .docx files)
    3. antiword/catdoc subprocess extraction (for binary .doc - requires system deps)
    4. Raw binary reading with common encodings

    Note: For binary .doc support, install one of:
      - Ubuntu: sudo apt-get install antiword catdoc
      - Mac: brew install antiword catdoc

    Args:
        file_path: Path to the .doc file

    Returns:
        Extracted text as string
    """
    # Try 1: python-docx (for .docx files mislabeled as .doc)
    try:
        doc = Document(str(file_path))
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception:
        pass

    # Try 2: docx2txt
    try:
        return docx2txt.process(file_path)
    except Exception:
        pass

    # Try 3: antiword/catdoc subprocess extraction for binary .doc
    for tool_name in ("antiword", "catdoc"):
        tool_path = shutil.which(tool_name)
        if not tool_path:
            continue

        try:
            result = subprocess.run(
                [tool_path, str(file_path)],
                capture_output=True,
                check=False,
                text=True,
                encoding="utf-8",
                errors="replace",
            )
        except Exception:
            continue

        if result.returncode == 0 and result.stdout.strip():
            logger.debug("Extracted %s with %s", file_path, tool_name)
            return result.stdout

    # Try 4: Raw binary reading with common encodings
    try:
        raw = file_path.read_bytes()
        # Try UTF-8, cp1251 (Russian), cp1252 (Windows)
        for encoding in ["utf-8", "cp1251", "cp1252", "iso-8859-1", "utf-16"]:
            try:
                return raw.decode(encoding, errors="replace")
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="replace")
    except Exception as e:
        raise ValueError(f"Failed to extract text from {file_path}: {e}") from e


def _extract_xls(file_path: Path) -> str:
    """Extract data from .xls file to markdown table.

    Args:
        file_path: Path to the .xls file

    Returns:
        Markdown table string
    """
    workbook = xlrd.open_workbook(file_path)
    sheet = workbook.sheet_by_index(0)

    md_lines = []

    # Header row
    headers = [str(cell.value).strip() for cell in sheet.row(0)]
    if headers:
        md_lines.append("|" + "|".join(headers) + "|")
        md_lines.append("|" + "|".join("-" for _ in headers) + "|")

    # Data rows
    for row_idx in range(1, sheet.nrows):
        row_data = [str(cell.value).strip() if cell.value else "" for cell in sheet.row(row_idx)]
        md_lines.append("|" + "|".join(row_data) + "|")

    return "\n".join(md_lines)


def convert_to_markdown(file_path: Path, overwrite: bool = True) -> tuple[bool, str | None, str | None, Path | None]:
    """Convert a single file to markdown based on its extension.

    Args:
        file_path: Path to the file to convert
        overwrite: If True, overwrite existing .md file (default: True)

    Returns:
        Tuple of (success, markdown_content_or_none, error_msg_or_none, md_path_or_none)
        - success: True if conversion succeeded and file was written
        - markdown_content: The extracted content (if successful)
        - error_msg: Error message (if failed) or special codes:
          - "IS_PDF": File is PDF, skip (handled separately)
          - "MD_ALREADY_EXISTS": MD file exists and overwrite=False
        - md_path: Path to the generated MD file (if successful), None otherwise
    """
    ext = file_path.suffix.lower()

    try:
        if ext == ".docx":
            content = _extract_docx(file_path)
        elif ext == ".doc":
            content = _extract_doc(file_path)
        elif ext == ".xls":
            content = _extract_xls(file_path)
        elif ext == ".pdf":
            # PDF handled separately via OCR
            return (False, None, "IS_PDF", None)
        else:
            return (False, None, f"Unsupported extension: {ext}", None)

        # Write MD file
        md_path = file_path.with_suffix(".md")

        if md_path.exists() and not overwrite:
            return (False, None, "MD_ALREADY_EXISTS", None)

        md_path.write_text(content, encoding="utf-8")
        logger.debug("Converted %s to %s", file_path, md_path)
        return (True, content, None, md_path)

    except Exception as e:
        error_msg = f"{type(e).__name__}: {e}"
        logger.error("Failed to convert %s: %s", file_path, error_msg)
        return (False, None, error_msg, None)


async def process_pdf_files(
    folder_path: Path,
    overwrite: bool = True,
) -> tuple[int, int, list[str], int, list[tuple[Path, Path]]]:
    """Process all PDF files in a folder with Mistral OCR.

    Reads PDF bytes and passes them to Mistral OCR as a base64 data URI.

    Args:
        folder_path: Path to folder containing PDF files
        overwrite: Whether to overwrite existing .md files (default: True)

    Returns:
        Tuple of (success_count, failure_count, failed_files, skipped_count, converted_pairs)
        where converted_pairs is a list of (pdf_path, md_path) tuples for successful conversions
    """
    pdf_files = [f for f in folder_path.iterdir() if f.is_file() and f.suffix.lower() == ".pdf"]

    if not pdf_files:
        return (0, 0, [], 0, [])

    results = []
    converted_pairs: list[tuple[Path, Path]] = []

    for pdf_path in pdf_files:
        try:
            md_path = pdf_path.with_suffix(".md")

            if md_path.exists() and not overwrite:
                results.append((False, "MD_ALREADY_EXISTS", str(pdf_path)))
                continue

            raw_bytes = await asyncio.to_thread(pdf_path.read_bytes)

            if len(raw_bytes) > MAX_PDF_SIZE_BYTES:
                results.append((False, f"PDF_TOO_LARGE({len(raw_bytes)} bytes)", str(pdf_path)))
                continue

            b64 = base64.b64encode(raw_bytes).decode("ascii")
            del raw_bytes

            data_uri = f"data:application/pdf;base64,{b64}"
            del b64

            request = OCRRequest(
                model="mistral-ocr-latest",
                document=DocumentURLChunk(document_url=data_uri, document_name=pdf_path.name),
            )
            result = await mistralai_ocr(request)

            ocr_text = "\n\n".join(page.markdown for page in result.pages)

            md_path.write_text(ocr_text, encoding="utf-8")
            logger.info("OCR converted %s to %s (%d chars)", pdf_path, md_path, len(ocr_text))
            results.append((True, None, str(pdf_path)))
            converted_pairs.append((pdf_path, md_path))

        except Exception as e:
            error_msg = f"{type(e).__name__}: {e}"
            logger.error("OCR failed for %s: %s", pdf_path, error_msg)
            results.append((False, error_msg, str(pdf_path)))

    success = sum(1 for suc, _, _ in results if suc)
    failure = sum(1 for suc, err, _ in results if not suc and err != "MD_ALREADY_EXISTS")
    skipped = sum(1 for _, err, _ in results if err == "MD_ALREADY_EXISTS")
    failed_files = [path for suc, err, path in results if not suc and err != "MD_ALREADY_EXISTS"]

    return (success, failure, failed_files, skipped, converted_pairs)


async def convert_company_files(
    company_name: str,
    folder_path: Path,
    semaphore: asyncio.Semaphore,
    overwrite: bool = True,
) -> tuple[str, int, int, list[str], list[tuple[Path, Path]]]:
    """Convert all eligible non-PDF/non-MD files in a company folder.

    Args:
        company_name: Company name for logging
        folder_path: Path to company folder
        semaphore: Concurrency limiter
        overwrite: Whether to overwrite existing .md files (default: True)

    Returns:
        Tuple of (company_name, success_count, failure_count, failed_files, converted_pairs)
        where converted_pairs is a list of (source_path, md_path) tuples for successful conversions
    """
    if not folder_path.exists():
        logger.warning("Company folder does not exist: %s", folder_path)
        return (company_name, 0, 0, [], [])

    # Find non-PDF and non-MD files only (MD files are already converted)
    files = [f for f in folder_path.iterdir() if f.is_file() and f.suffix.lower() not in (".pdf", ".md")]

    if not files:
        logger.debug("No non-PDF/non-MD files found for %s", company_name)
        return (company_name, 0, 0, [], [])

    logger.info("Found %d non-PDF/non-MD files to convert for %s", len(files), company_name)

    tasks = []
    for file_path in files:
        async with semaphore:
            task = asyncio.create_task(asyncio.to_thread(convert_to_markdown, file_path, overwrite))
            tasks.append(task)

    results = await asyncio.gather(*tasks)

    # Filter out PDF markers and track successful conversions
    converted_pairs: list[tuple[Path, Path]] = []
    success = 0
    failure = 0
    failed_files = []

    for i, file_path in enumerate(files):
        suc, _, err, md_path = results[i]
        if suc and err != "IS_PDF":
            success += 1
            if md_path:
                converted_pairs.append((file_path, md_path))
        elif not suc and err != "IS_PDF":
            failure += 1
            failed_files.append(str(file_path))

    logger.info(
        "Converted %d/%d non-PDF/non-MD files for %s (%d failed)",
        success,
        len(files),
        company_name,
        failure,
    )

    return (company_name, success, failure, failed_files, converted_pairs)


async def convert_all_files(
    results: list[tuple[str, list[str]]],
    output_root: Path,
    overwrite: bool = True,
    cleanup_source: bool = True,
) -> dict:
    """Convert all files for all companies to markdown.

    Handles both non-PDF files (using Python libraries) and PDF files
    (using Mistral OCR separately).

    Args:
        results: List of (company_name, urls) tuples
        output_root: Root output directory
        overwrite: Whether to overwrite existing .md files (default: True)
        cleanup_source: If True, remove source files (PDF, DOC, DOCX, XLS) after
            successful MD conversion. Failed files are kept. (default: True)

    Returns:
        Dictionary with conversion statistics:
        - total_companies
        - total_files_attempted
        - total_successful
        - total_failed
        - total_skipped
        - failed_files (list)
        - by_company (per-company breakdown)
        - cleaned_up_files (list) - NEW: list of source files removed
    """
    semaphore = asyncio.Semaphore(MAX_CONCURRENT_CONVERSIONS)

    # Process non-PDF files for all companies
    non_pdf_tasks = []
    pdf_folders = []

    for company_name, _ in results:
        folder_name = get_company_folder_name(company_name)
        folder_path = output_root / folder_name

        # Non-PDF conversion
        task = asyncio.create_task(convert_company_files(company_name, folder_path, semaphore, overwrite))
        non_pdf_tasks.append(task)
        pdf_folders.append((company_name, folder_path))

    # Run non-PDF conversions
    non_pdf_results = await asyncio.gather(*non_pdf_tasks)

    # Process PDF files for all companies (using Mistral OCR)
    pdf_tasks = []
    for _company_name, folder_path in pdf_folders:
        task = asyncio.create_task(process_pdf_files(folder_path, overwrite))
        pdf_tasks.append(task)

    pdf_results = await asyncio.gather(*pdf_tasks)

    # Aggregate stats and collect all converted pairs for cleanup
    stats = {
        "total_companies": len(results),
        "total_files_attempted": 0,
        "total_successful": 0,
        "total_failed": 0,
        "total_skipped": 0,
        "failed_files": [],
        "by_company": {},
        "cleaned_up_files": [],
    }

    all_converted_pairs: list[tuple[Path, Path]] = []

    # Aggregate non-PDF results
    for company, success, failure, failed_files, non_pdf_converted_pairs in non_pdf_results:
        stats["by_company"][company] = {
            "non_pdf_success": success,
            "non_pdf_failed": failure,
            "non_pdf_failed_files": failed_files,
        }
        stats["total_files_attempted"] += success + failure
        stats["total_successful"] += success
        stats["total_failed"] += failure
        stats["failed_files"].extend(failed_files)
        all_converted_pairs.extend(non_pdf_converted_pairs)

    # Aggregate PDF results
    for idx, (company_name, _) in enumerate(pdf_folders):
        success, failure, failed_files, skipped, pdf_converted_pairs = pdf_results[idx]
        if company_name not in stats["by_company"]:
            stats["by_company"][company_name] = {}
        stats["by_company"][company_name]["pdf_success"] = success
        stats["by_company"][company_name]["pdf_failed"] = failure
        stats["by_company"][company_name]["pdf_failed_files"] = failed_files
        stats["total_files_attempted"] += success + failure + skipped
        stats["total_successful"] += success
        stats["total_failed"] += failure
        stats["total_skipped"] += skipped
        stats["failed_files"].extend(failed_files)
        all_converted_pairs.extend(pdf_converted_pairs)

    # Cleanup source files after successful conversion
    if cleanup_source and all_converted_pairs:
        for source_path, md_path in all_converted_pairs:
            if md_path.exists():
                try:
                    source_path.unlink()
                    stats["cleaned_up_files"].append(str(source_path))
                    logger.info("Removed source file after conversion: %s", source_path)
                except Exception as e:
                    logger.warning("Failed to remove source file %s: %s", source_path, e)

    logger.info(
        "Conversion complete: %d attempted, %d successful, %d failed, %d skipped, %d source files cleaned up",
        stats["total_files_attempted"],
        stats["total_successful"],
        stats["total_failed"],
        stats["total_skipped"],
        len(stats["cleaned_up_files"]),
    )

    return stats
