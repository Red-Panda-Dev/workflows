"""Document conversion logic for EPFR workflow.

Converts company files (docx, doc, xls, xlsx) to Markdown format.
No PDF support - this workflow does not use OCR.
"""

import asyncio
import logging
from pathlib import Path
import shutil
import subprocess
import tempfile

from docx import Document
import docx2txt
import openpyxl
import xlrd

from .markdown_cleanup import clean_markdown_text


logger = logging.getLogger(__name__)

MAX_CONCURRENT_CONVERSIONS = 5


def _table_to_md(table) -> str:
    """Render a Word table as Markdown for disclosure text output.

    Args:
        table: ``python-docx`` table extracted from a disclosure document.

    Returns:
        Markdown table text preserving row and cell content.

    """
    normalized_rows: list[list[str]] = []

    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]

        while cells and not cells[0]:
            cells.pop(0)
        while cells and not cells[-1]:
            cells.pop()

        if cells:
            normalized_rows.append(cells)

    if not normalized_rows:
        return ""

    unique_widths = {len(row) for row in normalized_rows}
    if len(unique_widths) != 1 or next(iter(unique_widths)) == 1:
        return "\n".join(" | ".join(row) for row in normalized_rows)

    md_lines = ["|" + "|".join(normalized_rows[0]) + "|", "|" + "|".join("-" for _ in normalized_rows[0]) + "|"]
    for row in normalized_rows[1:]:
        md_lines.append("|" + "|".join(row) + "|")
    return "\n".join(md_lines)


def _extract_docx(file_path: Path) -> str:
    """Extract text and tables from a DOCX disclosure document.

    Args:
        file_path: Path to the downloaded DOCX file.

    Returns:
        Markdown-ready document text with tables converted to Markdown.

    """
    doc = Document(str(file_path))
    md_content = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            md_content.append(text)

    for table in doc.tables:
        md_table = _table_to_md(table)
        if md_table:
            md_content.append("")
            md_content.append(md_table)
            md_content.append("")

    return "\n\n".join(md_content)


def _extract_doc(file_path: Path) -> str:
    """Extract text from a legacy DOC disclosure document.

    Some EPFR records use misleading extensions, so the converter tries OOXML
    readers first, then xlrd (for OLE2 Excel files mislabelled as .doc),
    then binary DOC tools, then LibreOffice as a last resort.

    Args:
        file_path: Path to the downloaded DOC file.

    Returns:
        Extracted text suitable for writing to Markdown.

    Raises:
        ValueError: If no extraction strategy can read the file.

    """
    try:
        doc = Document(str(file_path))
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception:
        pass

    try:
        return docx2txt.process(file_path)
    except Exception:
        pass

    # Some files have a .doc extension but are actually OLE2 Excel workbooks.
    # xlrd handles both OLE2 (.xls) and mislabelled .doc Excel files.
    try:
        return _extract_xls(file_path)
    except Exception:
        pass

    for tool_name in ("antiword", "catdoc"):
        tool_path = shutil.which(tool_name)
        if not tool_path:
            continue

        try:
            result = subprocess.run(
                [tool_path, str(file_path)],
                capture_output=True,
                check=False,
                text=True,
                encoding="utf-8",
                errors="replace",
            )
        except Exception:
            continue

        if result.returncode == 0 and result.stdout.strip():
            logger.debug("Extracted %s with %s", file_path, tool_name)
            return result.stdout

    soffice_path = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice_path:
        try:
            with tempfile.TemporaryDirectory(prefix="epfr-doc-") as temp_dir:
                result = subprocess.run(
                    [
                        soffice_path,
                        "--headless",
                        "--convert-to",
                        "txt:Text",
                        "--outdir",
                        temp_dir,
                        str(file_path),
                    ],
                    capture_output=True,
                    check=False,
                    text=True,
                    encoding="utf-8",
                    errors="replace",
                )

                if result.returncode == 0:
                    txt_path = Path(temp_dir) / f"{file_path.stem}.txt"
                    if txt_path.exists():
                        text = txt_path.read_text(encoding="utf-8", errors="replace").strip()
                        if text:
                            logger.debug("Extracted %s with LibreOffice", file_path)
                            return text
        except Exception:
            pass

    raise ValueError(f"Failed to extract text from {file_path}")


def _extract_xls(file_path: Path) -> str:
    """Extract the first XLS worksheet as a Markdown table.

    Args:
        file_path: Path to the downloaded XLS file.

    Returns:
        Markdown table text from the first worksheet.

    """
    workbook = xlrd.open_workbook(file_path)
    sheet = workbook.sheet_by_index(0)

    md_lines = []

    if sheet.nrows > 0:
        headers = [str(cell.value).strip() for cell in sheet.row(0)]
        if headers:
            md_lines.append("|" + "|".join(headers) + "|")
            md_lines.append("|" + "|".join("-" for _ in headers) + "|")

    for row_idx in range(1, sheet.nrows):
        row_data = [str(cell.value).strip() if cell.value else "" for cell in sheet.row(row_idx)]
        md_lines.append("|" + "|".join(row_data) + "|")

    return "\n".join(md_lines)


def _extract_xlsx(file_path: Path) -> str:
    """Extract the first XLSX worksheet as a Markdown table.

    Args:
        file_path: Path to the downloaded XLSX file.

    Returns:
        Markdown table text from the active worksheet.

    """
    workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    sheet = workbook.active

    md_lines = []
    rows = list(sheet.iter_rows(values_only=True))

    if not rows:
        return ""

    headers = [str(cell).strip() if cell is not None else "" for cell in rows[0]]
    if headers:
        md_lines.append("|" + "|".join(headers) + "|")
        md_lines.append("|" + "|".join("-" for _ in headers) + "|")

    for row in rows[1:]:
        row_data = [str(cell).strip() if cell is not None else "" for cell in row]
        md_lines.append("|" + "|".join(row_data) + "|")

    workbook.close()
    return "\n".join(md_lines)


def convert_to_markdown(file_path: Path, overwrite: bool = True) -> tuple[bool, str | None, str | None, Path | None]:
    """Convert one supported EPFR office document to Markdown.

    Produces a sibling ``.md`` artifact for the mapping stage while preserving
    an explicit error string for unsupported or unreadable source files.

    Args:
        file_path: Path to the file to convert.
        overwrite: If True, overwrite an existing Markdown file.

    Returns:
        Tuple of success flag, Markdown content when conversion succeeds, error
        message when conversion fails, and Markdown path when written.

    """
    ext = file_path.suffix.lower()

    try:
        if ext == ".docx":
            content = _extract_docx(file_path)
        elif ext == ".doc":
            content = _extract_doc(file_path)
        elif ext == ".xls":
            content = _extract_xls(file_path)
        elif ext == ".xlsx":
            content = _extract_xlsx(file_path)
        else:
            return (False, None, f"Unsupported extension: {ext}", None)

        content = clean_markdown_text(content)

        # Write MD file
        md_path = file_path.with_suffix(".md")

        if md_path.exists() and not overwrite:
            return (False, None, "MD_ALREADY_EXISTS", None)

        md_path.write_text(content, encoding="utf-8")
        logger.debug("Converted %s to %s", file_path, md_path)
        return (True, content, None, md_path)

    except Exception as exc:
        error_msg = f"{type(exc).__name__}: {exc}"
        logger.error("Failed to convert %s: %s", file_path, error_msg)
        return (False, None, error_msg, None)


async def convert_unp_files(
    unp: str,
    folder_path: Path,
    semaphore: asyncio.Semaphore,
    overwrite: bool = True,
) -> tuple[str, int, int, list[str], list[tuple[Path, Path]]]:
    """Convert all supported office documents for one company UNP.

    Scans a company folder for downloaded DOC/DOCX/XLS/XLSX files and converts
    each one to Markdown so the final mapping can point to normalized text.

    Args:
        unp: Company tax identifier for logging and statistics.
        folder_path: Path to the company folder.
        semaphore: Concurrency limiter shared across UNP folders.
        overwrite: Whether to overwrite existing Markdown files.

    Returns:
        Tuple containing UNP, success count, failure count, failed file paths,
        and source-to-Markdown pairs for successful conversions.

    """
    if not folder_path.exists():
        logger.warning("UNP folder does not exist: %s", folder_path)
        return (unp, 0, 0, [], [])

    convertible_extensions = {".docx", ".doc", ".xls", ".xlsx"}
    files = [f for f in folder_path.iterdir() if f.is_file() and f.suffix.lower() in convertible_extensions]

    if not files:
        logger.debug("No convertible files found for %s", unp)
        return (unp, 0, 0, [], [])

    logger.info("Found %d convertible files for %s", len(files), unp)

    tasks = []
    for file_path in files:
        async with semaphore:
            task = asyncio.create_task(asyncio.to_thread(convert_to_markdown, file_path, overwrite))
            tasks.append(task)

    results = await asyncio.gather(*tasks)

    converted_pairs: list[tuple[Path, Path]] = []
    success = 0
    failure = 0
    failed_files = []

    for file_index, file_path in enumerate(files):
        suc, _, err, md_path = results[file_index]
        if suc:
            success += 1
            if md_path:
                converted_pairs.append((file_path, md_path))
        else:
            failure += 1
            failed_files.append(str(file_path))

    logger.info("Converted %d/%d files for %s (%d failed)", success, len(files), unp, failure)

    return (unp, success, failure, failed_files, converted_pairs)


async def convert_all_files(
    unp_folders: list[str],
    output_root: Path,
    overwrite: bool = True,
    cleanup_source: bool = True,
) -> dict:
    """Convert supported office documents across all company folders.

    Aggregates per-UNP conversion results and optionally removes source office
    documents after successful Markdown creation to keep mapping output clean.

    Args:
        unp_folders: Company UNP folder names to process.
        output_root: Root output directory containing UNP folders.
        overwrite: Whether to overwrite existing Markdown files.
        cleanup_source: Whether to delete source files after successful conversion.

    Returns:
        Conversion statistics with total counts, failed files, cleaned-up source
        files, and per-UNP details.

    """
    semaphore = asyncio.Semaphore(MAX_CONCURRENT_CONVERSIONS)

    tasks = []
    for unp in unp_folders:
        folder_path = output_root / unp
        task = asyncio.create_task(convert_unp_files(unp, folder_path, semaphore, overwrite))
        tasks.append(task)

    unp_results = await asyncio.gather(*tasks)

    stats: dict = {
        "total_unps": len(unp_results),
        "total_files_attempted": 0,
        "total_successful": 0,
        "total_failed": 0,
        "failed_files": [],
        "cleaned_up_files": [],
        "by_unp": {},
    }

    all_converted_pairs: list[tuple[Path, Path]] = []

    for unp, success, failure, failed_files, converted_pairs in unp_results:
        stats["by_unp"][unp] = {
            "success": success,
            "failed": failure,
            "failed_files": failed_files,
            "converted_pairs": [(str(src), str(md)) for src, md in converted_pairs],
        }
        stats["total_files_attempted"] += success + failure
        stats["total_successful"] += success
        stats["total_failed"] += failure
        stats["failed_files"].extend(failed_files)
        all_converted_pairs.extend(converted_pairs)

    if cleanup_source and all_converted_pairs:
        for source_path, md_path in all_converted_pairs:
            if md_path.exists():
                try:
                    source_path.unlink()
                    stats["cleaned_up_files"].append(str(source_path))
                    logger.info("Removed source file after conversion: %s", source_path)
                except Exception as exc:
                    logger.warning("Failed to remove source file %s: %s", source_path, exc)

    logger.info(
        "Conversion complete: %d attempted, %d successful, %d failed, %d source files cleaned up",
        stats["total_files_attempted"],
        stats["total_successful"],
        stats["total_failed"],
        len(stats["cleaned_up_files"]),
    )

    return stats
